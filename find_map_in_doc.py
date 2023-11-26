from docx import Document

doc = Document('map.docx')
last_map_name = ''
last_map_level = ''

with open('Map.txt','w') as mapWordsFile:
    for i,paragraph in enumerate(doc.paragraphs):
        if 'Territories' in paragraph.text:
            prev_paragraph = doc.paragraphs[i-1]
            preprev_paragraph = doc.paragraphs[i-2]
            if any(char.isdigit() for char in preprev_paragraph.text):
                last_map_level = preprev_paragraph.text[-1]
            map_name = prev_paragraph.text
            mapWordsFile.write(last_map_level + '\t')
            mapWordsFile.write(last_map_name + '\n')
            last_map_name = map_name
